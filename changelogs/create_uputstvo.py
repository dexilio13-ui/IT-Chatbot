#!/usr/bin/env python3
"""Skripta za kreiranje Uputstvo-za-prezentaciju.docx sa azuriranim kodom."""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "Uputstvo-za-prezentaciju.docx"

BLUE = RGBColor(0x1A, 0x56, 0xDB)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
DARK = RGBColor(0x33, 0x33, 0x33)


def set_run_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Calibri"
    return h


def add_label(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_code_block(doc, *lines):
    """Dodaj vise-linijski kod kao monospace blok sa sivom pozadinom."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = p.paragraph_format.element.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F2F2F2"}
    )
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK
    return p


def add_bullet(doc, text, bold_prefix="", size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, size=size)
        r2 = p.add_run(text)
        set_run_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size)
    return p


def add_mentor(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, italic=True, size=11)
    return p


def create_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- Naslovna strana ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("IT Asistent \u2014 RAG Chatbot")
    set_run_font(r, bold=True, size=26, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Uputstvo za prezentaciju sistema")
    set_run_font(r, size=16, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("FastAPI + LlamaIndex + OpenAI + Qdrant Cloud + React\nJul 2026.")
    set_run_font(r, size=11, color=LIGHT_GRAY)

    doc.add_page_break()

    # ========== Korak 1 ==========
    add_heading_styled(doc, "\U0001f680 Korak 1: Dolazak na aplikaciju i Login", level=1)

    add_label(doc, "Sta radis: ",
              "Otvaras React frontend aplikacije u pretrazivacu. Pojavljuje se LoginForm. "
              "Unosis kredencijale za Admina (admin / admin123).")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📝 Šta govoriš mentoru:")
    set_run_font(r, bold=True, color=BLUE)

    add_mentor(doc,
        '"Kada korisnik pristupi aplikaciji, prvo ga doceka React frontend i nasa Login forma. '
        "Ovde se vrsi autentifikacija. Ja cu se sada ulogovati kao Admin. U pozadini, frontend salje "
        "POST zahtev nasem FastAPI backendu sa mojim korisnickim imenom i lozinkom. Backend proverava "
        "bazu, hesira lozinku pomocu bcrypt-a i vraca mi JWT (JSON Web Token). Taj token se cuva u "
        "aplikaciji i Axios interceptor ga automatski 'lepi' u zaglavlje svakog sledeceg zahteva, "
        'kako bi sistem znao ko sam i koja su moja prava."')

    add_label(doc, "🔍 Šta da pokažeš u kodu (Ctrl+F): ", "")

    add_bullet(doc, "  -> pronadi rutu za login", bold_prefix="backend/api/auth.py")
    add_code_block(doc, '@router.post("/login")')
    add_bullet(doc, "  -> linija gde se kreira JWT token", bold_prefix="create_access_token")
    add_code_block(doc,
        "access_token = create_access_token(",
        "    data={'sub': user.username, 'role_id': user.role_id, 'is_admin': user.is_admin}",
        ")")

    # ========== Korak 2 ==========
    add_heading_styled(doc, "\U0001f4c1 Korak 2: Admin Panel i Upload Dokumenata (Parsiranje)", level=1)

    add_label(doc, "Sta radis: ",
              "Pokazujes mentoru kako Admin ubacuje dokumenta (PDF, DOCX) u sistem da bi baza "
              "znanja imala informacije.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📝 Šta govoriš mentoru:")
    set_run_font(r, bold=True, color=BLUE)

    add_mentor(doc,
        '"Posto sam Admin, imam prava da dodajem dokumente u bazu znanja. '
        "Ovde se desava kljucan proces \u2013 ekstrakcija teksta. Posto PDF fajlovi cesto imaju tabele "
        "i slozen format, obicni parseri tu prave haos. Zato mi koristimo LlamaParse alat. On uzima "
        "dokument i pretvara ga u cist Markdown format, cuvajuci strukturu tabela netaknutom. "
        "Zatim taj tekst ide kroz MarkdownNodeParser koji ga 'secka' (chunk-uje) na manje logicke "
        'celine pre slanja u bazu."')

    add_label(doc, "🔍 Šta da pokažeš u kodu (Ctrl+F): ", "")

    add_bullet(doc, "  -> pronadi upload endpoint i LlamaParse logiku", bold_prefix="backend/api/admin.py")
    add_code_block(doc, 'LlamaParse(result_type="markdown")')
    add_code_block(doc, "node_parser = MarkdownNodeParser()")

    # ========== Korak 3 ==========
    add_heading_styled(doc, "\U0001f9e0 Korak 3: Vektorizacija (Embeddings) i Qdrant Baza", level=1)

    add_label(doc, "Sta radis: ",
              "Nastavljas objasnjenje sta se desilo sa tim isecenim tekstom koji je upravo parsiran.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📝 Šta govoriš mentoru:")
    set_run_font(r, bold=True, color=BLUE)

    add_mentor(doc,
        '"Kada smo iseckali tekst, sistem ne moze tek tako da ga cita. Moramo da ga pretvorimo u brojeve, '
        "odnosno vektore, da bi vestacka inteligencija razumela kontekst. Za to koristimo OpenAI model "
        "'text-embedding-3-small'. On pretvara svaki chunk u vektor od 1536 dimenzija. Te vektore onda "
        "cuvamo u Qdrant Cloud vektorskoj bazi. Vazno je naglasiti: uz svaki vektor, u bazu upisujemo i "
        "metapodatak 'required_role_id'. To nam omogucava strogu RBAC kontrolu \u2013 dokument za servisere "
        'dobija id 3, i Qdrant ga strogo cuva."')

    add_label(doc, "🔍 Šta da pokažeš u kodu (Ctrl+F): ", "")

    add_bullet(doc, "  -> pronadi embedding i vektorsku konfiguraciju",
               bold_prefix="backend/rag/engine.py i backend/core/config.py")
    add_code_block(doc, 'Settings.embed_model = OpenAIEmbedding(model="text-embedding-3-small")')
    add_code_block(doc,
        "from llama_index.vector_stores.qdrant import QdrantVectorStore",
        "QdrantVectorStore(client=_client, collection_name=settings.QDRANT_COLLECTION)")

    # ========== Korak 4 ==========
    add_heading_styled(doc, "\U0001f4ac Korak 4: Postavljanje pitanja (RAG i RBAC mehanizam)", level=1)

    add_label(doc, "Sta radis: ",
              'Prelazis na chat interfejs i kucas pitanje, na primer: "Koje je resenje za gresku ESRV0102?"')

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📝 Šta govoriš mentoru:")
    set_run_font(r, bold=True, color=BLUE)

    add_mentor(doc,
        '"Sada korisnik postavlja pitanje u chatu. Ovde nas RAG (Retrieval-Augmented Generation) sistem '
        "stupa na snagu. Pre nego sto bilo sta posaljemo LLM-u, moramo da nadjemo relevantne podatke iz "
        "baze. FastAPI backend dekodira moj JWT token, vidi koja je moja uloga (role_id) i salje upit u "
        "Qdrant. Ono sto je kljucno za bezbednost: filter se aplicira na nivou vektorske pretrage! "
        "Qdrant ce mi vratiti samo one vektore ciji je 'required_role_id' manji ili jednak mojoj ulozi. "
        "Ako kupac (role 1) pita za serviserski kod (role 3), vektorska baza taj podatak fizicki blokira "
        'i on nikada ne stigne do AI modela."')

    add_label(doc, "🔍 Šta da pokažeš u kodu (Ctrl+F): ", "")

    add_bullet(doc, "  -> pronadi RBAC filter koji se primenjuje na vektorsku pretragu",
               bold_prefix="backend/rag/engine.py")
    add_code_block(doc,
        "filters = MetadataFilters(",
        "    filters=[",
        '        MetadataFilter(',
        '            key="required_role_id",',
        "            value=role_id,",
        "            operator=FilterOperator.LTE,",
        "        )",
        "    ]",
        ")")

    # ========== Korak 5 ==========
    add_heading_styled(doc, "\U0001f916 Korak 5: LLM, Sistemski Prompt i Generisanje Odgovora", level=1)

    add_label(doc, "Sta radis: ",
              "Dobijas tacan odgovor na ekranu sa prikazanim izvorima (Citations) iz prirucnika.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📝 Šta govoriš mentoru:")
    set_run_font(r, bold=True, color=BLUE)

    add_mentor(doc,
        '"Kada Qdrant pronadje dozvoljene i najslicnije tekstove (kontekst), to pakujemo zajedno sa '
        "korisnickim pitanjem i nasim strogim Sistemskim Promptom, i tek onda saljemo glavnom LLM-u. "
        "Koristimo 'gpt-4o-mini' sa jako niskom temperaturom od 0.1 da bismo sprecili mastanje i "
        "halucinacije. Sve to objedinjuje 'CondensePlusContextChatEngine', koji je najmocniji engine u "
        "LlamaIndex-u, jer on prvo kondenzuje istoriju mog razgovora da bi razumeo kontekst, a onda "
        'generise finalni odgovor iskljucivo na osnovu ubacenih dokumenata."')

    add_label(doc, "🔍 Šta da pokažeš u kodu (Ctrl+F): ", "")

    add_bullet(doc, "  -> pronadi LLM konfiguraciju", bold_prefix="backend/rag/engine.py")
    add_code_block(doc, 'Settings.llm = OpenAI(model="gpt-4o-mini", temperature=0.1)')

    add_bullet(doc, "  -> pronadi chat engine", bold_prefix="CondensePlusContextChatEngine")
    add_code_block(doc,
        "return CondensePlusContextChatEngine.from_defaults(",
        "    retriever=hybrid_retriever,",
        "    memory=memory,",
        "    system_prompt=SYSTEM_PROMPT,",
        "    llm=Settings.llm,",
        ")")

    add_bullet(doc, " (opciono, ako pita za prevenciju halucinacija) -> sistemski prompt",
               bold_prefix="backend/rag/system_prompt.py")
    add_code_block(doc,
        "1. Odgovaraj ISKLJUCIVO na osnovu prilozenog konteksta iz baze znanja.",
        "2. Ako odgovor nije eksplicitno sadrzan u kontekstu, odgovori tacno:",
        "   'Nemam taj podatak u bazi znanja.'")

    # ========== Rezime ==========
    doc.add_page_break()
    add_heading_styled(doc, "Rezime za prezentaciju", level=1)

    summary = (
        "Kada budes izlagao, drzi se ovog toka. Logika je linearna:\n\n"
        "   Korisnik (Frontend) -> Token (Auth) -> Upload (Parser) -> "
        "Vektori (Embedding + Qdrant) -> Pitanje (RBAC Filter) -> Odgovor (LLM)\n\n"
        'Ako te pita "Zasto bas taj alat?", uvek imas spreman inzenjerski '
        "razlog u recenicama iznad. Srecno na prezentaciji!"
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(summary)
    set_run_font(r, size=12)

    # Tabela sa test nalozima
    add_heading_styled(doc, "Test nalozi", level=2)
    table = doc.add_table(rows=5, cols=5)
    table.style = "Light Grid Accent 1"

    headers = ["Username", "Password", "role_id", "Uloga", "is_admin"]
    data = [
        ["admin",     "admin123", "3", "Admin",      "\u2705"],
        ["serviser",  "123",      "3", "Technician", "\u274c"],
        ["prodavac",  "123",      "2", "Sales",      "\u274c"],
        ["kupac",     "123",      "1", "Customer",   "\u274c"],
    ]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for par in table.rows[0].cells[i].paragraphs:
            for run in par.runs:
                run.bold = True

    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            table.rows[ri + 1].cells[ci].text = ct

    doc.save(str(OUTPUT_PATH))
    print(f"Dokument kreiran: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_document()
